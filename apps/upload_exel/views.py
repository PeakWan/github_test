import base64
import json
import os
import pickle
import re
from django.contrib.auth import authenticate, logout
from django.shortcuts import render, redirect
from django.urls import reverse
import matplotlib.pyplot as plt
import random
import time
import threading
# import xlwt
import eventlet

import uuid
import pyarrow as pa

from io import StringIO, BytesIO
import pandas as pd
from django.http import FileResponse
from aliyunsdkcore.http.format_type import JSON
# from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render
from numpy.matlib import randn
from docx import Document
from docx.shared import Inches

from AnalysisFunction.X_1_DataGovernance import miss_data_delete, miss_data_filling, get_var_low_colinear, \
    get_var_low_vif, data_balance, dummies_recoding, psm_matching, group_recoding, data_standardization, \
    non_numerical_value_process, abnormal_deviation_process, _feature_classification, data_manipulate, data_transform
from AnalysisFunction.X_2_DataSmartStatistics import normal_test, two_sample_t_test
from AnalysisFunction.X_3_DataSeniorStatistics import survival_estimating_models
from AnalysisFunction.X_3_R_DataSeniorStatistic import R_surv_ana

from apps.index_qt.models import File_old, Browsing_process, Member, MemberType
from django.http import HttpResponse, JsonResponse
from django import http
from django.views import View
from django_redis import get_redis_connection
from django.db.models import Max
from datetime import datetime
from utils.member_verification import member_verification
from libs.get import get_s, loop_add, write, read, get_path, filtering_view
from rest_framework.views import APIView
from apps.index_qt.jwt_token import MyBaseAuthentication

class NpEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, pd.np.integer):
            return int(obj)
        elif isinstance(obj, pd.np.floating):
            return float(obj)
        elif isinstance(obj, pd.np.ndarray):
            return obj.tolist()
        else:
            return super(NpEncoder, self).default(obj)


pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', None)

#authentication_classes = [MyBaseAuthentication, ]
class Normality(APIView):
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        """
        实现正态性校验
        :param request:
        :return:
        """
        context = get_s(request, project_id, data_id)
        # 单选加多选
        context["value"] = "much"
        context["begin"] = "begin"
        return render(request, 'index/normality.html', context=context)

    def post(self, request, project_id, data_id):
        """
        接收传的列名
        :param request:
        :return:
        """

        json_data = json.loads(request.body.decode())
        electedText = json_data.get("name")
        num = json_data.get("number")
        decimal_num = json_data.get("num_select")
        # 数据筛选
        filter_data = json_data.get("filter")
        # 使用要分析的列名

        if not all([electedText, str(num)]):
            return HttpResponse(json.dumps({'code': '201', 'error': '请选择变量'}))

        # 创建一个uuid唯一值
        id = uuid.uuid1()
        id = str(id)
        # 获取传的列名
        if int(num) == 0:
            # 获取数据库的对象
            user = File_old.objects.get(id=project_id)
            # 获取原文件的路径
            old_file_path = user.path_pa

        elif int(num) == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            old_file_path = re_d.get("df_result")
        try:
            # 打开文件
            # 查询该用户的用户等级
            try:
                grade = Member.objects.get(user_id=request.user.id)
                # 查询当前用户的等级
                member = MemberType.objects.get(id=grade.member_type_id)
            except Exception as e:
                return http.JsonResponse({'code': 1002, 'error': '会员查询失败，请稍后重试'})
            df_r = read(old_file_path)
            df = df_r.iloc[0:int(member.number)]
           # 写入数据库的路径
            end01 = "/static/"
            end = "/"
            # 本地代码上面
            # end = "\\"
            if int(num) == 0:
                string3 = old_file_path[:old_file_path.rfind(end)]
            else:
                string = old_file_path[:old_file_path.rfind(end)]
                string3 = string[:string.rfind(end)]
            # strint3路径 为用户创的 项目下面 

            # string2 为图片的静态路径
            string_tp = string3[string3.rfind(end01):]

            # file_s 为这方法存图片以及df表的文件夹  绝对路径
            file_s = string3 + "/" + "正态性校验" + id
            try:
                # 创建存图片和df表的文件夹的文件夹
                os.makedirs(file_s)
            except Exception as e:
                print(e)
                return HttpResponse(json.dumps({'code': '1003', 'error': '创建失败'}, cls=NpEncoder))
            
            savepath = file_s + '/'

            if filter_data:
                df = filtering_view(df,json_data)
            # 调用正态性验证方法
            try:  # df_input, continuous_features,decimal_num=3
                
                re = normal_test(df, continuous_features=electedText, decimal_num=int(decimal_num),savePath=savepath)
            except Exception as e:
                print(e)
                return http.JsonResponse({'code': 1004, 'error': '分析失败，请选择正确变量'})
            try:
                if re['error']:
                    return http.JsonResponse({'code': 1005, 'error': re['error']})
            except Exception as e:
                print('正常')
            # 提取plt  和str 数据
            df_result = re[0]
            plts = re[2]
            str_s = re[1]
            str_s = str_s.replace("\n", "<br/>")
            # 获取返还给页面的df数据
            a = df_result
            before = {}
            before["name"] = "结果数据描述"
            if list(a):
                # 保存第一列的列名
                before['Listing'] = []
                name = list(a)
                # 循环列名
                for i in range(len(list(a))):
                    before['Listing'].append(name[i])
                # 添加数据
                before['info'] = []
                for n in range(len((a[name[1]]))):
                    dict_x = {}
                    y = 0
                    # 制定第一行 所有列的所有info数据
                    for i in name:
                        i = str(i)
                        # 最多添加八个数据
                        # if y == 8:
                        #     break
                        try:
                            if dict_x[i]:
                                t = str(uuid.uuid1())
                                i = str(i) + t
                                dict_x[i] = str(a.iloc[n][y])
                        except Exception as e:
                            dict_x[i] = str(a.iloc[n][y])
                        y += 1
                    before['info'].append(dict_x)
 
                # 把得到的图片保存
            # 返回前端 图片的静态路径
            pltlist = loop_add(plts, "正态性校验", string_tp, id)

            # plt.close()
            # 写入数据
            df_file = file_s + '/' + 'df_result.pkl'
            write(df_result, df_file)
            browsing_process = {}
            browsing_process['name'] = '正态性校验'
            browsing_process['str_result'] = str_s
            browsing_process['df_result_2'] = df_file
            for i in range(len(pltlist)):
                if i == 0:
                    name = 'plt'
                    browsing_process[name] = pltlist[i]
                else:
                    name = 'plt' + str(i + 1)
                    browsing_process[name] = pltlist[i]
            try:
                process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                    Max('order'))
                print(process)
                if process:
                    order = int(process['order__max']) + 1
                    print(order)
                    browsing = Browsing_process.objects.create(process_info=browsing_process,
                                                               user_id=request.user.id,
                                                               order=str(order),
                                                               file_old_id=project_id)
                    project = browsing.id

            except Exception as e:
                print(e)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order='1',
                                                           file_old_id=project_id)
                project = browsing.id

            # 把返回结果数据封装到列表里面
            form = [before]
            # img = [plt_file, plt_file2]
            context = {
                'project': project,
                'img': pltlist,
                'str': str_s,
                "form": form,
                "text": "tableDescriptionImg"

            }
        except Exception as e:
            print(e)
            return HttpResponse(json.dumps({'code': '201', 'error': '网络延迟，请稍后再试'}))
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': context}, cls=NpEncoder))


class StatisticView(APIView):
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        """
        t检验页面
        :param request:
        :return:
        """
        context = get_s(request, project_id, data_id)
        # 单选加多选
        context["value"] = "All"
        context["begin"] = "begin"
        context["name"] = "tjy"
        return render(request, 'index/statistic2.html', context=context)

    def post(self, request, project_id, data_id):
        # 定义数据 因为前端未写好先自己定义数据
        json_data = json.loads(request.body.decode())
        # 获取传过来的第一个列名
        y_column = json_data.get("y_column")
        # 获取传过来的连续数字
        group_labels = json_data.get("y_labels")
        # 判断用户是原始数据还是分析后的数据
        num = json_data.get("number")
        # 判断用户传的类型
        related = json_data.get("relate")

        #  获取小数点位
        decimal_num = json_data.get("num_select")
        continuous_list = json_data.get("arr")
        # 数据筛选
        filter_data = json_data.get("filter")
        if related == "False":
            related = False
        elif related == "True":
            related = True
        if not group_labels:
            group_labels = None
        else:
            try:
                group_labels = list(map(int, group_labels))
            except Exception as e:
                print(e)
        # 创建一个uuid唯一值
        id = uuid.uuid1()
        id = str(id)

        # 获取传的列名
        if num == 0:
            # 使用要分析的列名
            electedText = y_column
            # 获取数据库的对象
            user = File_old.objects.get(id=project_id)
            # 获取原文件的路径
            file_path = user.path_pa

        elif num == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 打开文件
            # 查询该用户的用户等级
            try:
                grade = Member.objects.get(user_id=request.user.id)
                # 查询当前用户的等级
                member = MemberType.objects.get(id=grade.member_type_id)
            except Exception as e:
                return http.JsonResponse({'code': 1002, 'error': '会员查询失败，请重试'})
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
            
            if filter_data:
                df = filtering_view(df,json_data)
            try:
                re = two_sample_t_test(df, group=y_column, continuous_features=continuous_list,
                                       group_labels=group_labels, relate=related, decimal_num=int(decimal_num))
            except Exception as e:
                print(e)
                return http.JsonResponse({'code': 1005, 'error': '分析失败，请选择正确变量'})
            try:
                if re['error']:
                    return http.JsonResponse({'code': 1005, 'error': re['error']})
            except Exception as e:
                print('正常')
            # 提取df  和str 数据
            plts = re[0]
            a = plts
            before = {}
            before["name"] = "结果数据描述"
            if list(a):
                # 保存第一列的列名
                before['Listing'] = []
                name = list(a)
                print(name, len(name), "-=-=-=-=-=")
                # 循环列名
                # print(len(list(a)) + 1)
                for i in range(len(list(a))):
                    before['Listing'].append(name[i])
                    # if i == 8:
                    #     break
                # 添加数据
                before['info'] = []
                for n in range(len((a[name[1]]))):
                    # dict_x = {'name': a.iloc[n].name}
                    dict_x = {}
                    y = 0
                    # 制定第一行 所有列的所有info数据
                    for i in name:
                        i = str(i)
                        try:
                            if dict_x[i]:
                                t = str(uuid.uuid1())
                                i = str(i) + t
                                dict_x[i] = str(a.iloc[n][y])
                        except Exception as e:
                            dict_x[i] = str(a.iloc[n][y])
                        y += 1
                    before['info'].append(dict_x)

            str_s = re[1]
            # print(str_s, "=====")
            str_s = str_s.replace("\n", "<br/>")
            # 服务器上面
            end = "/"
            # 本地代码上面
            # end = "/"

            if num == 0:
                string2 = file_path[:file_path.rfind(end)]
            else:
                string = file_path[:file_path.rfind(end)]
                string2 = string[:string.rfind(end)]
            # print(string2)
            try:
                # 创建一个唯一文件夹
                os.mkdir(string2 + '/t检验' + id)
                # print(string2, "123123123")
            except Exception as e:
                print(e)
                return HttpResponse(json.dumps({'code': '1003', 'error': '创建失败'}, cls=NpEncoder))
            # 写入数据
            df_file = string2 + '/t检验' + id + '/' + 'df_result.pkl'
            write(plts, df_file)
            browsing_process = {}
            browsing_process['name'] = 't检验'
            browsing_process['df_result_2'] = df_file
            browsing_process['str_result'] = str_s

            try:
                process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                    Max('order'))
                print(process)
                # print(process)
                if process:
                    order = int(process['order__max']) + 1
                    print(order)
                    browsing = Browsing_process.objects.create(process_info=browsing_process,
                                                               user_id=request.user.id,
                                                               order=order,
                                                               file_old_id=project_id)
                    project = browsing.id

            except Exception as e:
                print(e)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=1,
                                                           file_old_id=project_id)
                project = browsing.id

            # 把返回结果数据封装到列表里面
            form = [before]
            context = {
                'project': project,
                'str': str_s,
                "form": form,
                "text": "tableDescription",
            }
        except Exception as e:
            print(e)
            return HttpResponse(json.dumps({'code': '201', 'error': '网络延迟，请稍后再试！'}))
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': context}, cls=NpEncoder))



class UploadView(APIView):
    """上传文件的类"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request):
        """
        接收上传的文件
        :param request:
        :return:
        """
        # 获取参数
        a = request.FILES.get('upfile')
        project_name = request.POST.get('name')
        project_background = request.POST.get('background')
        project_outline = request.POST.get('outline')
        print(project_name)
        # 校验参数
        if not project_name:
            return http.JsonResponse({'code': 1001, 'error': '请填写项目名称'})
        if not project_background:
            return http.JsonResponse({'code': 1002, 'error': '请填写项目背景'})
        if not a:
            return http.JsonResponse({'code': 1003, 'error': '请上传所分析文件'})
        # 查询该用户的用户等级
        grade = Member.objects.get(user_id=request.user.id)
        # 查询当前用户的等级
        member = MemberType.objects.get(id=grade.member_type_id)
        # 查询当前用户已经创建项目的个数
        files = File_old.objects.filter(user_id=request.user.id).all()
        if int(member.projects) <= len(files):
            return http.JsonResponse({'code': 10010, 'error': '会员等级不足，创建失败'})
        try:
            filename = a.name
            b = filename.split('.')
            m = '.'+b[-1]
            # print(b[-1])
            # print('----------------------------')
            # m = re.findall(r'\..+', filename)
            # print(m[0])
            # 判断文件名后缀是否被允许打开
            if m == '.csv' or m == '.xlsx' or m == '.xls':
                # 将最新上传的文件保存
                # 保存的路径
                # 用户项目文件夹
                x = uuid.uuid1()
                # print(x)
                project_path=get_path(0)  # 得上一层路径  /www/wwwroot/znfxpt.ceshi.vip/ini/Intelligent
                user_file_path = 'user' + str(request.user.id) + '_' + project_name + str(x) + '/'
                os.makedirs(str(project_path) + '/static' + '/exel/' + './user' + str(
                    request.user.id) + '_' + project_name + str(x))
                path = str(project_path)+'/static' + '/exel/' + user_file_path + filename
                with open(path, 'wb+') as f:
                    # print(a.chunks(), "=============================")
                    for chunk in a.chunks():
                        f.write(chunk)
                # 读取文件
                try:
                    try:
                        f = pd.read_excel(path)
                    except Exception as e:
                        print(e)
                        f = pd.read_csv(path, encoding='GBK')
                except Exception as e:
                    print(e)
                    f = pd.read_csv(path)
                # 获取表的列数
                df_num = f.shape[0] # projects
                
                str_re = "您上传的数据总共包含%s例样本,目前您的会员等级%s,分析最大样本量为%s例。" % (df_num,member.member_name, member.number)
                print("2342342343253454367546756876575698")

                print(str_re)
                a = f.columns
                list01 = []
                try:
                    # test_case = ['h el.lo', '...', 'h3.a', 'ds_4,']
                    #         def characters_replace(old_list,characters_list,characters_target):
                    #             for c in characters_list:
                    #                 old_list=[x.replace(c,characters_target) for x in old_list]
                    #             return old_list
                            
                    #         print(characters_replace(test_case,['.',' '],'_'))
                    for i in a:
                        print(i)
                        c = re.sub(u"([^\u4e00-\u9fa5\u0030-\u0039\u0041-\u005a\u0061-\u007a,])", "", str(i))
                        if c[0:1].isnumeric():
                            c = "_" + c
                        if c not in list01:
                            list01.append(c)
                        else:
                            return JsonResponse({'code': 10040, 'error': '表格数据内有重复列名，请修改后上传'})
                except Exception as e:
                    print(e)
                    return JsonResponse({'code': 10030, 'error': '表格数据有非法字符'})
                f.columns = list01
                # 重新写入数据
                try:
                    f.to_excel(path, index=False)
                except Exception as e:
                    print(e)
                    f.to_csv(path, index=False)
                column_name = list(f)
                c = filename.split('.')
                path_pa = str(project_path) +'/static' + '/exel/' + user_file_path + str(
                    c[0]) + '.pkl'
                write(f, path_pa)
                # 将文件名字保存到数据库
                # print(11111111111111111)
                try:
                    user = request.user
                    last_time = str(datetime.now())
                    File_old.objects.create(file_name=filename, user_id=user.id, project_name=project_name,
                                            background=project_background, outline=project_outline, path=path,
                                            last_time=last_time, path_pa=path_pa)
                    user.dft_file = path
                    user.save()
                except Exception as e:
                    print(e)
                    return JsonResponse({'code': 1010, 'error': '保存数据库失败'})
                # print(2222222222222222222222222222222222222222)
                # 查询项目数据表
                data = []
                info = File_old.objects.all()
                # print(33333333333333333333333333333)
                # 循环所有的项目
                for num in info:
                    i = {}
                    i['id'] = num.id
                    i['project_name'] = num.project_name
                    i['background'] = num.background
                    i['outline'] = num.outline
                    i['file_name'] = num.file_name
                    i['user_id'] = num.user.id
                    data.append(i)
            # 不允许上传的数据类型
            else:
                print('文件上传失败')
                return http.JsonResponse({'code': '2002', 'error': '文件上传失败,只允许上传XLS\XLSX\CSV数据'})
        # 抛出错误正常执行下面的代码
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': '2003', 'error': '文件上传失败,只允许上传XLS\XLSX\CSV数据'})
            

        if int(df_num) < int(member.number):
            str_re = None
        # print(1111111111111)
        # print(str_re)
        context = {
            'column': column_name,
            'data': data,
            'str_re': str_re,
        }
        # 返回页面
        return http.JsonResponse({'code': '200', 'error': '添加成功', 'context': context})


class Index_tou(APIView):
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id):
        # 获取会员图标
        try:
            member_num = Member.objects.get(user=request.user.id)
        except Exception as e:
            print(e)
            return render(request, 'index/index_sy.html')
        u_id = member_num.member_type.id
        # 非法访问
        try:
            project_name = File_old.objects.get(id=project_id, user_id=request.user.id)
            print(project_name)
        except Exception as e:
            print(1111111)
            print(e)
            # 清理session
            logout(request)
        
            # 退出登录，重定向到登录页
            response = redirect(reverse('index_qt:index_sy'))
        
            # 退出登录时清除cookie中的username
            response.delete_cookie('username')
            response.delete_cookie('uuid')
            response.delete_cookie('token')
        
            return response

        context = {
            'project_id': project_id,
            'u_id': u_id,

        }

        return render(request, 'index/home.html', context=context)


class MissDeleteView(APIView):
    """缺失值开始分析"""
    
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id):
        context = {
            'value': 'None',
            'project_id': project_id
        }
        return render(request, 'index/miss_delete.html', context=context)

    def post(self, request, project_id):
        """
        缺乏值异常处理
        :param request:
        :return:
        """
        # 接收参数
        json_data = json.loads(request.body.decode())
        # 获取参数
        select = json_data.get('selevt')
        missing_rate = json_data.get('missing_rate')
        number = json_data.get('number')
        # 校验参数
        if not all([select, missing_rate]):
            return http.HttpResponseForbidden('变量不能为空')
        # 获取传的列名
        if number == 0:
            # 获取数据库的对象
            user = File_old.objects.get(id=project_id, user_id=request.user.id)
            # 获取原文件的路径
            old_file_path = user.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            old_file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            # 读取dataframe数据
            df_r = read(old_file_path)
            # 查看会员级别以及读取多少条
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '请登录后重新上传数据'})
        # 读取要分析的数据
        try:
            info = miss_data_delete(df_input=df, miss_rate=float(missing_rate), miss_axis=int(select))
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1003, 'error': '分析失败，请选择正确的变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 分析文字描述
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")

        # 原始分类数据描述
        c = info[2]
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): str(c.iloc[num].name)})
                    else:
                        input_o_desc['info'][num][str(i)] = str(c.iloc[num][i - 1])

        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): str(a.iloc[num].name)})
                    else:
                        before['info'][num][str(i)] = str(a.iloc[num][i - 1])

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): str(d.iloc[num].name)})
                    else:
                        result_o_desc['info'][num][str(i)] = str(d.iloc[num][i - 1])

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): str(b.iloc[num].name)})
                    else:
                        after['info'][num][str(i)] = str(b.iloc[num][i - 1])

        # 将浏览过程添加到数据库
        end = "/"
        string2 = None
        if number == 0:
            old_file_path = old_file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = old_file_path[:old_file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/miss_delete' + str(select) + '_' + str(missing_rate))
        except Exception as e:
            print(e)

        # 写入数据
        write(info[0], string2 + '/miss_delete' + str(select) + '_' + str(missing_rate) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/miss_delete' + str(select) + '_' + str(missing_rate) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/miss_delete' + str(select) + '_' + str(missing_rate) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/miss_delete' + str(select) + '_' + str(missing_rate) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/miss_delete' + str(select) + '_' + str(missing_rate) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '缺失数据删除'
        browsing_process['df_result'] = string2 + '/miss_delete' + str(select) + '_' + str(
            missing_rate) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/miss_delete' + str(select) + '_' + str(
            missing_rate) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/miss_delete' + str(select) + '_' + str(
            missing_rate) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/miss_delete' + str(select) + '_' + str(
            missing_rate) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/miss_delete' + str(select) + '_' + str(
            missing_rate) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        form = [before, after, input_o_desc, result_o_desc]
        data = {
            "form": form,
            'project': project,
            'str_result': str_s
        }

        # 返回页面
        return http.JsonResponse({'code': '200', 'error': '分析成功', 'context': data})
        # return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class SmartView(APIView):
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id):
        """
        智能文章生成页面
        :param request:
        :return:
        """
        data = {}
        step = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id, is_delete='1').all()
        for i in step:
            info = i.process_info
            info2 = eval(info)
            info2['id'] = i.id
            data[i.id] = info2
        context = {
            'data': data,
            'project_id': project_id,
            'begin': 'article'
        }
        return render(request, 'index/Smart_article.html', context=context)

    def post(self, request, project_id):
        data = {}
        step = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id, is_delete='1').all()
        for i in step:
            info = i.process_info
            info2 = eval(info)
            info2['id'] = i.id
            data[i.id] = info2
            if not info2.get("str_result", 0):
                info2["str_result"] = "分析结果无"
                
            df_result_2 = info2.get("df_result_2", 0)
            df_result_2_2 = info2.get("df_result_2_2", 0)
            df_result_2_2_2 = info2.get("df_result_2_2_2", 0)

            # print(df_result_2, "-=-=-=-=-=")
            if df_result_2:
                df_dict = {}
                # 打开文件
                # 查询该用户的用户等级
                df = read(df_result_2)
                try:
                    df_re = df.fillna('')
                except Exception as e:
                    print(e)
                    df_re = df
                # 计量数据描述
                # print(list(df_re), "=========list(df_re)=====")
                # print(df_re)
                if list(df_re):
                    d = df_re
                    # 保存第一列的列名
                    df_dict['Listing'] = []
                    # 所有列名的列表
                    name = list(d)
                    # del (name[0])
                    # 循环列名
                    for i in range(len(name)):
                        df_dict['Listing'].append(name[i])
                    # 添加数据
                    df_dict['info'] = []
                    
                    # print("===131241421========")
                    
                    # print(name, "-===========")
                    # print(type(name), "-===========")
                    # # print()
                    # print(len(name[0]), "-===-=-=-=")
                    
                    for num in range(len(d[name[0]])):
                        dict_x = {}
                        y = 0
                        # 制定第一行 所有列的所有info数据
                        for i in name:
                            i = str(i)
                            # print(dict_x)
                            # print(d.iloc[0][1], "-==-=-=11111")
                            # print(num, "---num--", y, "-y-")
                            dict_x[i] = str(d.iloc[num][y])
                            y += 1
                        df_dict['info'].append(dict_x)

                    info2["df_result_2"] = df_dict

            if df_result_2_2:
                df_dict = {}
                # 打开文件
                # 查询该用户的用户等级
                # try:
                #     grade = Member.objects.get(user_id=request.user.id)
                #     # 查询当前用户的等级
                #     member = MemberType.objects.get(id=grade.member_type_id)            
                # except Exception as e:
                #     return http.JsonResponse({'code': 1002, 'error': '会员查询失败，请重试'})
                df = read(df_result_2_2)
                # df = df_r.iloc[0:int(member.number)]
                # try:
                #     try:
                #         df = pd.read_excel(df_result_2_2, nrows=int(member.number))
                #     except Exception as e:
                #         print(e)
                #         df = pd.read_csv(df_result_2_2, nrows=int(member.number))
                # except Exception as e:
                #     print(e)
                #     df = pd.read_csv(df_result_2_2, nrows=int(member.number), encoding='gbk')  
                try:
                    d = df.fillna('')
                except Exception as e:
                    print(e)
                    d = df
                # 计量数据描述
                if list(d):
                    # 保存第一列的列名
                    df_dict['Listing'] = []
                    # 所有列名的列表
                    name = list(d)
                    # del (name[0])
                    # 循环列名
                    for i in range(len(name)):
                        df_dict['Listing'].append(name[i])
                    # 添加数据
                    df_dict['info'] = []
                    for num in range(len(d[name[0]])):
                        dict_x = {}
                        y = 0
                        # 制定第一行 所有列的所有info数据
                        for i in name:
                            i = str(i)
                            dict_x[i] = str(d.iloc[num][y])
                            y += 1
                        df_dict['info'].append(dict_x)

                    info2["df_result_2_2"] = df_dict

            if df_result_2_2_2:
                df_dict = {}
                # 打开文件
                # 查询该用户的用户等级
                # try:
                #     grade = Member.objects.get(user_id=request.user.id)
                #     # 查询当前用户的等级
                #     member = MemberType.objects.get(id=grade.member_type_id)            
                # except Exception as e:
                #     return http.JsonResponse({'code': 1002, 'error': '会员查询失败，请重试'})
                df = read(df_result_2_2_2)
                # df = df_r.iloc[0:int(member.number)]
                # try:
                #     try:
                #         df = pd.read_excel(df_result_2_2_2, nrows=int(member.number))
                #     except Exception as e:
                #         print(e)
                #         df = pd.read_csv(df_result_2_2_2, nrows=int(member.number))
                # except Exception as e:
                #     print(e)
                #     df = pd.read_csv(df_result_2_2_2, nrows=int(member.number), encoding='gbk')    
                try:
                    d = df.fillna('')
                except Exception as e:
                    print(e)
                    d = df
                # 计量数据描述
                if list(d):
                    # 保存第一列的列名
                    df_dict['Listing'] = []
                    # 所有列名的列表
                    name = list(d)
                    # del (name[0])
                    # 循环列名
                    for i in range(len(name)):
                        df_dict['Listing'].append(name[i])
                    # 添加数据
                    df_dict['info'] = []
                    for num in range(len(d[name[0]])):
                        dict_x = {}
                        y = 0
                        # 制定第一行 所有列的所有info数据
                        for i in name:
                            i = str(i)
                            dict_x[i] = str(d.iloc[num][y])
                            y += 1
                        df_dict['info'].append(dict_x)

                    info2["df_result_2_2_2"] = df_dict
        print(data)
        context = {
            'data': data,
            'project_id': project_id,
        }
        # 返回页面
        # return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))
        return http.JsonResponse({'code': '200', 'error': '分析成功', 'context': data})
#  http.JsonResponse({'code': 200, 'error': '查询成功',"context":context})

class SaveView(APIView):
    """保存结果显示到智能"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request):
        # 获取参数
        json_dict = json.loads(request.body.decode())
        project_id = json_dict.get('project')
        user = request.user
        # get memberInfo By UserId
        member_info = Member.objects.get(user__id=user.id)
        flow_number_max = MemberType.objects.get(id=member_info.member_type_id).flow_number
        # user flow numbers count
        project_file_id = Browsing_process.objects.get(id=project_id).file_old_id
        flow_number = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_file_id,
                                                      is_delete='1').all().count()
        if (flow_number >= flow_number_max):
            return HttpResponse(json.dumps({'code': 1003, 'error': '已超出最大的流程保存数量,请尝试提升会员等级或删除其他部分保存的流程'}))        
        
        browsing = Browsing_process.objects.get(id=project_id)
        file_old_id = browsing.file_old_id
        browsing.is_delete = str(1)
        # 保存最新的结果
        re_d = eval(browsing.process_info)
        try:
            if re_d.get("df_result"):
                try:
                    # 找到这个项目的最新处理的数据
                    latest_data = Browsing_process.objects.get(file_old=browsing.file_old, is_latest=1)
                    if latest_data:
                        # 修改数据结果
                        latest_data.is_latest = 0
                        latest_data.save()
                except Exception as e:
                    print(e)
                browsing.is_latest = 1
        except Exception as e:
            print(e)
        browsing.save()
        # 获取当前保存的路径
        user = File_old.objects.get(id=file_old_id)
        path = user.path_pa
        # 切割路径
        end = '/'
        route01 = path[:path.rfind(end)]
        # 生成word文档的路径
        word_path = route01 + '/' + '文章结果.docx'
        # 创建word对象
        try:
            document = Document(word_path)
        except Exception as e:
            print(e)
            document = Document()
        # 添加步骤标题
        p_total = document.add_heading("")  # 添加标题
        r_total = p_total.add_run("步骤：" + re_d.get('name'))
        r_total.add_break()
        project_path=get_path(0)  # 得上一层路径  /www/wwwroot/znfxpt.ceshi.vip/ini/Intelligent
        # 添加字符串
        if re_d.get('str_result'):
            r_total01 = p_total.add_run(str(re_d.get('str_result')))
            r_total01.add_break()
        
        # 向文档里面添加图片
        if re_d.get('plt'):
            p_total02 = document.add_heading("图片描述：")  # 添加标题
            r_total02 = p_total02.add_run()
            r_total02.add_break()
            img_name = re_d.get('plt')
            image = str(project_path) + img_name
            document.add_picture(image, width=Inches(1.5))  # 向文档里添加图片
        if re_d.get('plt2'):
            img_name02 = re_d.get('plt2')
            image02 = str(project_path)  + img_name02
            document.add_picture(image02, width=Inches(1.5))  # 
        if re_d.get('plt3'):
            img_name03 = re_d.get('plt3')
            image03 = str(project_path)  + img_name03
            document.add_picture(image03, width=Inches(1.5))  # 
        if re_d.get('plt4'):
            img_name04 = re_d.get('plt4')
            image04 = str(project_path)  + img_name04
            document.add_picture(image04, width=Inches(1.5))  # 
        if re_d.get('plt5'):
            img_name05 = re_d.get('plt5')
            image05 = str(project_path)  + img_name05
            document.add_picture(image05, width=Inches(1.5))  # 
        name = None
        # 原始分类数据描述表格
        try:
            if re_d.get('df_result_2'):
                input01 = document.add_heading("")  # 添加标题
                input_o_desc = input01.add_run('结果描述表格:')
                input_o_desc.add_break()
                file_path = re_d.get('df_result_2')
                # 读取服务器上的表格
                df = read(file_path)
                # try:
                #     try:
                #         df = pd.read_excel(file_path)
                #     except Exception as e:
                #         print(e)
                #         df = pd.read_csv(file_path)
                # except Exception as e:
                #     print(e)
                #     df = pd.read_csv(file_path,encoding='gbk')
                list01 = list(df)
                name = list01
                print(name)
                info = []
                for num in range(len(df[name[0]])):
                    for i in range(len(name) + 1):
                        if i == 0:
                            info.append({str(i): df.iloc[num].name})
                        else:
                            info[num][str(i)] = df.iloc[num][i - 1]
                print(info)
                # # # 添加表格
                table = document.add_table(0, len(name) + 1, style="Medium Grid 1 Accent 1")
                for i in range(len(df[name[1]]) + 1):
                    row_cells = table.add_row().cells
                    if i == 0:
                        for b in range(len(name) + 1):
                            if b == 0:
                                row_cells[0].text = ''
                            else:
                                row_cells[b].text = name[b - 1]
                    else:
                        for c in range(len(name) + 1):
                            row_cells[c].text = str(info[i - 1][str(c)])
        except Exception as e:
            print(e)
        
        try:
            # 原始分类数据描述表格
            if re_d.get('df_result_2_2'):
                input02 = document.add_heading("")  # 添加标题
                input_o_desc02 = input02.add_run('结果描述表格第二个:')
                input_o_desc02.add_break()
                file_path = re_d.get('df_result_2_2')
                # 读取服务器上的表格
                df = read(file_path)
                list01 = list(df)
                name02 = list01
                print(name)
                info02 = []
                for num in range(len(df[name02[0]])):
                    for i in range(len(name02) + 1):
                        if i == 0:
                            info02.append({str(i): df.iloc[num].name})
                        else:
                            info02[num][str(i)] = df.iloc[num][i - 1]
                # # # 添加表格
                table = document.add_table(len(name), len(name) + len(name02) + 1, style="Medium Grid 1 Accent 1")
                for i in range(len(df[name02[1]]) + 1):
                    row_cells02 = table.add_row().cells
                    if i == 0:
                        for b in range(len(name02) + 1):
                            if b == 0:
                                row_cells02[0].text = ''
                            else:
                                row_cells02[b].text = name02[b - 1]
                    else:
                        for c in range(len(name02) + 1):
                            row_cells02[c].text = str(info02[i - 1][str(c)])
        except Exception as e:
            print(e)
        # 保存表格
        document.save(word_path)  # 保存文档

        return HttpResponse(json.dumps({'code': '200', 'error': '结果保存成功'}))


class MissDataFilling(APIView):
    """
    数据填补的函数
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):

        context = {
            # "data": data,
            'project_id': project_id,
            'value': 'much'
        }
        return render(request, 'index/miss_data_filling.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取需要处理的列名
        features = json_data.get('features')
        # 获取填补方法
        method = json_data.get('method')
        method = str(method)
        if method != 'KNN':
            method = method.lower()
        # 填补常数
        constant = json_data.get('constant')
        # 获取哪种数据
        number = json_data.get('number')
        # 校验参数
        if not all([features, method]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        if method == 'constant':
            if not constant:
                return HttpResponse(json.dumps({'code': 1003, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '请登录后重新上传数据后再操作'})
        # 将浏览过程添加到数据库
        end = "/"
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        uuid01 = uuid.uuid1()
        try:
            os.mkdir(string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method))
        except Exception as e:
            print(e)
        # 保存图片
        end01 = "/static/"
        route01 = string2[string2.rfind(end01):]
        route02 = route01 + '/miss_data_filling' + str(uuid01)
        path = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/'
        # 读取数据
        try:
            if constant:
                constant = float(constant)
                info = miss_data_filling(df, features, method=method, constant=constant,path=path)
            else:
                info = miss_data_filling(df, features, method=method,path=path)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1003, 'error': '分析失败，请选择正确的变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        try:
            # 原始分类数据描述
            c = info[2]
            input_o_desc = {}
            input_o_desc['name'] = '原始分类数据描述'
            if list(c):
                # 保存第一列的列名
                input_o_desc['Listing'] = []
                # 所有列名的列表
                name = list(c)
                # 循环列名
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['Listing'].append('名字')
                    else:
                        input_o_desc['Listing'].append(name[i - 1])
                # 添加数据
                input_o_desc['info'] = []
                for num in range(len(c['总数'])):
                    for i in range(len(list(c)) + 1):
                        if i == 0:
                            input_o_desc['info'].append({str(i): c.iloc[num].name})
                        else:
                            input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
            # 原始计量数据描述
            a = info[3]
            before = {}
            before['name'] = '原始计量数据描述'
            if list(a):
                # 保存第一列的列名
                before['Listing'] = []
                name = list(a)
                # 循环列名
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['Listing'].append('名字')
                    else:
                        before['Listing'].append(name[i - 1])
                # 添加数据
                before['info'] = []
                for num in range(len(a['总数'])):
                    for i in range(len(list(a)) + 1):
                        if i == 0:
                            before['info'].append({str(i): a.iloc[num].name})
                        else:
                            before['info'][num][str(i)] = a.iloc[num][i - 1]

            # 结果分类数据描述
            d = info[4]
            result_o_desc = {}
            result_o_desc['name'] = '结果分类数据描述'
            if list(d):
                # 保存第一列的列名
                result_o_desc['Listing'] = []
                # 所有列名的列表
                name = list(d)
                # 循环列名
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['Listing'].append('名字')
                    else:
                        result_o_desc['Listing'].append(name[i - 1])
                # 添加数据
                result_o_desc['info'] = []
                for num in range(len(d['总数'])):
                    for i in range(len(list(d)) + 1):
                        if i == 0:
                            result_o_desc['info'].append({str(i): d.iloc[num].name})
                        else:
                            result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

            # 结果计量数据描述
            b = info[5]
            after = {}
            after['name'] = '结果计量数据描述'
            if list(b):
                # 保存第一列的列名
                after['Listing'] = []
                # 所有列名的列表
                name = list(b)
                # 循环列名
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['Listing'].append('名字')
                    else:
                        after['Listing'].append(name[i - 1])
                # 添加数据
                after['info'] = []
                for num in range(len(b['总数'])):
                    for i in range(len(list(b)) + 1):
                        if i == 0:
                            after['info'].append({str(i): b.iloc[num].name})
                        else:
                            after['info'][num][str(i)] = b.iloc[num][i - 1]
        except Exception as e:
            return http.JsonResponse({'code': 1003, 'error': info})
        plts = info[6]
        print(plts)
        # 保存图片
        pltlist = loop_add(plts, "miss_data_filling", route01, str(uuid01)+ '_' + str(method))
        # 写入数据
        write(info[0], string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/miss_data_filling' + str(uuid01) + '_' + str(method) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '数据填补'
        browsing_process['df_result'] = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(
            method) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(
            method) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(
            method) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(
            method) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/miss_data_filling' + str(uuid01) + '_' + str(
            method) + '/' + 'result_n_desc.pkl'
        # 如果有图片保存到数据库
        if len(features) == 1:
            for i in range(len(pltlist)):
                if i == 0:
                    name = 'plt'
                    browsing_process[name] = pltlist[i]
                else:
                    name = 'plt' + str(i + 1)
                    browsing_process[name] = pltlist[i]
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s,
            'plts':pltlist
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class GetVarView(APIView):
    """
    获取相关性过高变量
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id):

        context = {
            'project_id': project_id
        }
        context["value"] = "much"
        return render(request, 'index/get_var_colinear.html', context=context)

    def post(self, request, project_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取需要处理的阀值
        thres = json_data.get('thres')
        thres = float(thres)
        # 获取哪种数据
        number = json_data.get('number')
        # 获取多选的列名
        features = json_data.get('name')
        print(features)
        if features == []:
            features = None
        # 校验参数
        if not all([thres]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]

        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '请登录后重新上传数据后再操作'})
        # 读取数据
        try:
            info = get_var_low_colinear(df, thres=thres, features=features)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1001, 'error': '分析失败，请选择正确的变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/get_var_low_colinear' + str(thres))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/get_var_low_colinear' + str(thres) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/get_var_low_colinear' + str(thres) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/get_var_low_colinear' + str(thres) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/get_var_low_colinear' + str(thres) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/get_var_low_colinear' + str(thres) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '获取相关性过高变量'
        browsing_process['df_result'] = string2 + '/get_var_low_colinear' + str(thres) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/get_var_low_colinear' + str(thres) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/get_var_low_colinear' + str(thres) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/get_var_low_colinear' + str(
            thres) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/get_var_low_colinear' + str(
            thres) + '' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            # 'result_o_desc': result_o_desc,
            # 'input_o_desc': input_o_desc,
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class GetVifView(APIView):
    """
    获取共线性过高变量
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id):
        context = {
            'project_id': project_id
        }
        context["value"] = "much"
        return render(request, 'index/get_var_vif.html', context=context)

    def post(self, request, project_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取需要处理的阀值
        thres = json_data.get('thres')
        thres = int(thres)
        # 获取哪种数据
        number = json_data.get('number')
        # 获取多选的列名
        features = json_data.get('name')
        print(features)
        if features == []:
            features = None
        # 校验参数
        if not all([thres]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            # 查询最新的原始数据
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '请登录后重新上传数据后再操作'})

        # 读取数据
        try:
            if thres:
                info = get_var_low_vif(df_input=df, thres=thres, features=features)
            elif not thres:
                info = get_var_low_vif(df_input=df, features=features)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        c = info[2]
        if list(c):
            list02 = ['count', 'categorical', 'highest', 'frequency']
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
                # 添加数据
                input_o_desc['info'] = []
                for num in range(len(c['总数'])):
                    for i in range(len(list(c)) + 1):
                        if i == 0:
                            input_o_desc['info'].append({str(i): c.iloc[num].name})
                        else:
                            input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]

                            # 原始计量数据描述
        before = {}
        a = info[3]
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
                # 添加数据
                before['info'] = []
                for num in range(len(a['总数'])):
                    for i in range(len(list(a)) + 1):
                        if i == 0:
                            before['info'].append({str(i): a.iloc[num].name})
                        else:
                            before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        d = info[4]
        if list(d):
            list02 = ['count', 'categorical', 'highest', 'frequency']
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]
        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/get_var_low_vif' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '获取共线性过高变量'
        browsing_process['df_result'] = string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/get_var_low_vif' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class DataBalanceView(APIView):
    """
    数据平衡
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):

        context = {
            'value': 'single',
            # "data": data,
            'project_id': project_id,
        }
        return render(request, 'index/data_balance.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取分组变量
        group_name = json_data.get('name')
        # 获取比例
        ratio = json_data.get('ratio')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取方式
        method = json_data.get('method')

        try:
            ratio01 = {}
            for k, v in ratio.items():
                ratio01[int(k)] = int(v)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1001, 'error': '填写的比例数据不正确'})

        # 校验参数
        if not all([group_name, ratio, method]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            # 查询最新的原始数据
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = data_balance(df, group_name, ratio01, method)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        # 校验是不是错误信息
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            print(len(c['总数']))
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]
        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/data_balance' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/data_balance' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/data_balance' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/data_balance' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/data_balance' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/data_balance' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '数据平衡'
        browsing_process['df_result'] = string2 + '/data_balance' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/data_balance' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/data_balance' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/data_balance' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/data_balance' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class PsmMatchinView(APIView):
    """
    PSM倾向性匹配
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):

        context = {
            'value': 'All',
            'project_id': project_id,
            'name': 'psm'
        }
        return render(request, 'index/psm_matchin.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取匹配的列表
        x_field = json_data.get('x_field')
        # 获取分组变量
        y_column = json_data.get('y_column')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取比例
        ratio = json_data.get('ratio')
        # 获取是否精准匹配
        precious = json_data.get('precious')

        precious_bool = None
        if precious == 'false':
            precious_bool = False
        else:
            precious_bool = True
        
        try:
            ratio = int(ratio)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '填写的比例不正确，请重新填写'})

            # 校验参数
        if not all([x_field, y_column, ratio]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 校验第二个字符串不能在第一个列表里面
        if y_column in x_field:
            return HttpResponse(json.dumps({'code': 1002, 'error': '单选框内不能与多选内的重复'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理最新的数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]

        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = psm_matching(df, list(x_field), y_column, ratio,precious_bool)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
                # if i == 8:
                #     break
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
                # if i == 8:
                #     break
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                print(num)
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/psm_matchin' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/psm_matchin' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/psm_matchin' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/psm_matchin' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/psm_matchin' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/psm_matchin' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = 'PSM倾向性匹配'
        browsing_process['df_result'] = string2 + '/psm_matchin' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/psm_matchin' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/psm_matchin' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/psm_matchin' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/psm_matchin' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class DummiesView(APIView):
    """
    哑变量重编码
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'single'
        return render(request, 'index/dummies_recoding.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取匹配的列表
        feature = json_data.get('feature')
        # 获取哪种数据
        number = json_data.get('number')

        # 校验参数
        if not all([feature]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = dummies_recoding(df_input=df, features=feature)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/dummies_recoding' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/dummies_recoding' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/dummies_recoding' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/dummies_recoding' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/dummies_recoding' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/dummies_recoding' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '哑变量重编码'
        browsing_process['df_result'] = string2 + '/dummies_recoding' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/dummies_recoding' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/dummies_recoding' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/dummies_recoding' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/dummies_recoding' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class GroupRecodingView(APIView):
    """
    分组重编码
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'single'
        return render(request, 'index/group_recoding.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取匹配的列表
        feature = json_data.get('feature')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取按什么来分组
        name = json_data.get('name')
        # 获取范围重编码
        value = json_data.get('value')
        # 获取数据类型
        value_type = json_data.get('type')
        # 校验参数
        if not all([feature]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]

        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            if name == 'groupnum':
                value = int(value)
                print(value_type)
                info = group_recoding(df_input=df, feature=feature, group_num=value, type=int(value_type))
            elif name == 'group_percentile':
                value = value.split(',')
                print(value)
                int01 = []
                for i in range(len(value) + 2):
                    try:
                        if i == 0:
                            int01.append(0)
                        else:
                            int01.append(float(value[i - 1]))
                    except Exception as e:
                        print(e)
                        int01.append(1)

                print(int01)
                info = group_recoding(df_input=df, feature=feature, group_percentile=int01, type=int(value_type))
            elif name == 'group_cut_value':
                value = value.split(',')
                int02 = []
                # 获取这一列数据的最大值和最小值
                max_min = list(df[feature])
                first_min = min(max_min)
                last_max = max(max_min)
                for i in range(len(value) + 2):
                    try:
                        if i == 0:
                            int02.append(first_min)
                        else:
                            int02.append(float(value[i - 1]))
                    except Exception as e:
                        print(e)
                        int02.append(last_max)
                
                info = group_recoding(df_input=df, feature=feature, group_cut_value=int02, type=int(value_type))
            elif name == 'group_cut_value2':
                # value = eval(value)
                print(value)
                new_value = {}
                # print(int(value))
                for x, y in value.items():
                    # print(x, y, "-=-=-=-=")
                    list01 = []
                    for a in y:
                        #list01.append(int(a))
                        list01.append(float(a))
                    for b in list01:
                        print(111111111)
                        print(type(b))
                        print(22222222222)
                    new_value[int(x)] = list01

                    # print(type(x))
                    # print(type(y))
                    # value[x] = value.pop(int(x))
                print(new_value, "-=-=-=new_value-=-=-=-")
                info = group_recoding(df_input=df, feature=feature, group_range_value=value, type=int(value_type))
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()

        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/group_recoding' + str(uuid01))
        except Exception as e:
            print(e)

        # 写入数据
        write(info[0], string2 + '/group_recoding' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/group_recoding' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/group_recoding' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/group_recoding' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/group_recoding' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '分组重编码'
        browsing_process['df_result'] = string2 + '/group_recoding' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/group_recoding' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/group_recoding' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/group_recoding' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/group_recoding' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        
        form = [before, after, input_o_desc, result_o_desc]
        # print(before)
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class DataStandardizationView(APIView):
    """
    数据标准化
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'much'
        return render(request, 'index/data_standardization.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取方式
        method = json_data.get('method')

        # 校验参数
        if not all([features, method]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = data_standardization(df_input=df, features=features, method=method)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        try:
            # 原始分类数据描述
            c = info[2]
            list02 = ['count', 'categorical', 'highest', 'frequency']
            input_o_desc = {}
            input_o_desc['name'] = '原始分类数据描述'
            if list(c):
                # 保存第一列的列名
                input_o_desc['Listing'] = []
                # 所有列名的列表
                name = list(c)
                # 循环列名
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['Listing'].append('名字')
                    else:
                        input_o_desc['Listing'].append(name[i - 1])
                # 添加数据
                input_o_desc['info'] = []
                for num in range(len(c['总数'])):
                    for i in range(len(list(c)) + 1):
                        if i == 0:
                            input_o_desc['info'].append({str(i): c.iloc[num].name})
                        else:
                            input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
            # 原始计量数据描述
            a = info[3]
            before = {}
            before['name'] = '原始计量数据描述'
            if list(a):
                # 保存第一列的列名
                before['Listing'] = []
                name = list(a)
                print(name)
                list01 = ['count', 'unique', 'top', 'freq']
                # 循环列名
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['Listing'].append('名字')
                    else:
                        before['Listing'].append(name[i - 1])
                # 添加数据
                before['info'] = []
                for num in range(len(a['总数'])):
                    for i in range(len(list(a)) + 1):
                        if i == 0:
                            before['info'].append({str(i): a.iloc[num].name})
                        else:
                            before['info'][num][str(i)] = a.iloc[num][i - 1]

            # 结果分类数据描述
            d = info[4]
            list02 = ['count', 'categorical', 'highest', 'frequency']
            result_o_desc = {}
            result_o_desc['name'] = '结果分类数据描述'
            if list(d):
                # 保存第一列的列名
                result_o_desc['Listing'] = []
                # 所有列名的列表
                name = list(d)
                # 循环列名
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['Listing'].append('名字')
                    else:
                        result_o_desc['Listing'].append(name[i - 1])
                # 添加数据
                result_o_desc['info'] = []
                for num in range(len(d['总数'])):
                    for i in range(len(list(d)) + 1):
                        if i == 0:
                            result_o_desc['info'].append({str(i): d.iloc[num].name})
                        else:
                            result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

            # 结果计量数据描述
            b = info[5]
            after = {}
            after['name'] = '结果计量数据描述'
            if list(b):
                # 保存第一列的列名
                after['Listing'] = []
                # 所有列名的列表
                name = list(b)
                # 循环列名
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['Listing'].append('名字')
                    else:
                        after['Listing'].append(name[i - 1])
                # 添加数据
                after['info'] = []
                for num in range(len(b['总数'])):
                    for i in range(len(list(b)) + 1):
                        if i == 0:
                            after['info'].append({str(i): b.iloc[num].name})
                        else:
                            after['info'][num][str(i)] = b.iloc[num][i - 1]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': info})

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/data_standardization' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/data_standardization' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/data_standardization' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/data_standardization' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/data_standardization' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/data_standardization' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '数据标准化'
        browsing_process['df_result'] = string2 + '/data_standardization' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/data_standardization' + str(
            uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/data_standardization' + str(
            uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/data_standardization' + str(
            uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/data_standardization' + str(
            uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class NonNumericalView(APIView):
    """
    非数值异常值处理
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        return render(request, 'index/non_numerical_value_process.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取方式
        method = json_data.get('method')

        # 校验参数
        if not all([features, method]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            # 查询最新的原始数据
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = non_numerical_value_process(df, features, method)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(i): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(i)] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(i): a.iloc[num].name})
                    else:
                        before['info'][num][str(i)] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(i): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(i)] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(i): b.iloc[num].name})
                    else:
                        after['info'][num][str(i)] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        print(number)
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/non_numerical_value_process' + str(uuid01))
        except Exception as e:
            print(e)
        # 写入数据
        write(info[0], string2 + '/non_numerical_value_process' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/non_numerical_value_process' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/non_numerical_value_process' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/non_numerical_value_process' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/non_numerical_value_process' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '非数值异常值处理'
        browsing_process['df_result'] = string2 + '/non_numerical_value_process' + str(
            uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/non_numerical_value_process' + str(
            uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/non_numerical_value_process' + str(
            uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/non_numerical_value_process' + str(
            uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/non_numerical_value_process' + str(
            uuid01) + '/' + 'result_n_desc.pkl'
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }
        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class AbnormalView(APIView):
    """
    异常偏离值处理
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'much'
        return render(request, 'index/abnormal_deviation_process.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取方式
        method = json_data.get('method')
        # 偏离比例float
        ratio = json_data.get('ratio')

        # 校验参数
        if not all([features, method]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        end = "/"
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/abnormal_deviation_process' + str(uuid01))
        except Exception as e:
            print(e)

        # plts.plot(randn(10).cumsum(), 'k--')
        # 设置图片大小
        end01 = "/static/"
        route01 = string2[string2.rfind(end01):]
        route02 = route01 + '/abnormal_deviation_process' + str(uuid01)
        path = string2 + '/abnormal_deviation_process' + str(uuid01) + '/'
        try:
            if ratio:
                print(11111111111111)
                ratio = float(ratio)
                info = abnormal_deviation_process(df_input=df, features=features, method=method, ratio=ratio, path=path)
            else:
                print(22222222222222222222)
                info = abnormal_deviation_process(df_input=df, features=features, method=method, path=path, ratio=1.5)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 保存图片并展示
        plts = info[6]
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(c.iloc[num].name): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(list(c)[i - 1])] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(a.iloc[num].name): a.iloc[num].name})
                    else:
                        before['info'][num][str(list(a)[i - 1])] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(d.iloc[num].name): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(list(d)[i - 1])] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(b.iloc[num].name): b.iloc[num].name})
                    else:
                        after['info'][num][str(list(b)[i - 1])] = b.iloc[num][i - 1]
        # 将浏览过程添加到数据库
        pltlist = loop_add(plts, "abnormal_deviation_process", route01, str(uuid01))
        # 写入数据
        write(info[0], string2 + '/abnormal_deviation_process' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/abnormal_deviation_process' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/abnormal_deviation_process' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/abnormal_deviation_process' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/abnormal_deviation_process' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '异常偏离值处理'
        browsing_process['df_result'] = string2 + '/abnormal_deviation_process' + str(
            uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/abnormal_deviation_process' + str(
            uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/abnormal_deviation_process' + str(
            uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/abnormal_deviation_process' + str(
            uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/abnormal_deviation_process' + str(
            uuid01) + '/' + 'result_n_desc.pkl'
        # 如果有图片保存到数据库
        if len(features) == 1:
            for i in range(len(pltlist)):
                if i == 0:
                    name = 'plt'
                    browsing_process[name] = pltlist[i]
                else:
                    name = 'plt' + str(i + 1)
                    browsing_process[name] = pltlist[i]
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        # if len(features) == 1:
        data = {
            'form': form,
            'project': project,
            'str_result': str_s,
            'plt': pltlist[0],
            'plt2': pltlist[1]
        }
        # else:
        #     data = {
        #     'form': form,
        #     'project': project,
        #     'str_result':str_s
        #     }

        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class GroupingView(APIView):
    """获取列的数据类型以及数量"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project_id):
        """查询相关列的信息"""
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 校验参数
        if not all([features]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if int(number) == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif int(number) == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]

        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1005, 'error': '上传文件有误请重新上传'})
        # 读取列的数据
        name = list(df[str(features)])
        num = {}
        # 计算数据的类型和数量
        for i in name:
            if i not in num:
                num[i] = 1
            else:
                b = num[i] + 1
                num[i] = b
        print(num)
        num01 = []
        print(num.items())
        for k, v in num.items():
            b = {}
            b['key'] = k
            b['value'] = v
            num01.append(b)

        # 判断数据类型是否大于5
        if len(list(num.keys())) > 5:
            return http.JsonResponse({'code': 1006, 'error': '该列数据类型大于5列请重新选择'})

        # 返回页面
        return http.JsonResponse({'code': 200, 'error': '分析成功', 'context': num01})


class ListingView(APIView):
    """获取列名的类"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request):
        """获取列名"""
        # 定义数据 因为前端未写好先自己定义数据
        json_data = json.loads(request.body.decode())
        # 获取传过来的第一个列名
        num = json_data.get("list")
        project_id = json_data.get('project_id')
        # 判断是原始数据还是处理后的数据
        # 获取传的列名
        if num == 0:
            # 获取数据库的对象
            print(project_id, "====project_id===")
            user = File_old.objects.get(id=project_id)
            # 获取原文件的路径
            old_file_path = user.path_pa
            # old_file_path = user.path
        elif num == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
                re_d = eval(process.process_info)
                old_file_path = re_d.get("df_result")
            except Exception as e:
                print(e)
                return http.JsonResponse({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'})
        # 打开文件
        df = read(old_file_path)

        listing = list(df)
        # print(df.dtypes)
        ret = _feature_classification(df)
        new_list = {}
        for i, y in enumerate(ret):
            if y:
                if i == 0:
                    new_list["continuous_features"] = y
                elif i == 1:
                    new_list["categorical_features"] = y
                elif i == 2:
                    new_list["time_features"] = y

        context = {
            'listing': listing,
            'new_list': new_list,

        }
        return http.JsonResponse({'code': 200, 'error': '查询成功', "context": context})


class ButtonView(APIView):
    """获取数据有没有处理过的数据"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request):
        # 定义数据 因为前端未写好先自己定义数据
        json_data = json.loads(request.body.decode())
        # 获取当前项目的id
        project_id = json_data.get('project_id')

        # 查询当前项目有没有处理后的数据
        # 查询所有之前的顺序
        button_id = None
        # 查询最新的原始数据
        try:
            process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
            if file_path:
                button_id = 1
        except Exception as e:
            print(e)
            button_id = 0

        context = {
            'button_id': button_id
        }
        return http.JsonResponse({'code': 200, 'error': '查询成功', "context": context})


class Download(APIView):
    """下载处理后文件到本地"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project):
        # 校验参数
        print(project)
        if not all([project]):
            return http.JsonResponse({'code': 1001, 'error': '导出失败，缺失参数'})
        # 查找最终的结果
        try:
            info = Browsing_process.objects.get(file_old_id=project, is_latest=1,user_id=request.user.id)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1003, 'error': '没有最新的数据表格'})
        info = eval(info.process_info)
        # 获取最新路径
        if info.get('df_result'):
            path = info.get('df_result')
            # 先读取dateframe数据
            df_r = read(path)
            # 写入数据表中

            write_path = path[:path.rfind('/')]
            path = write_path + '/df_result.xlsx'
            df_r.to_excel(path, index=False)
        else:
            return http.JsonResponse({'code': 1002, 'error': '没有保存最新的处理数据'})
        # 切割路径
        end = '/static/'
        route01 = path[path.rfind(end):]
        context = {
            'path': route01
        }
        return http.JsonResponse({'code': 200, 'error': '成功', 'context': context})


class Orifinal(APIView):
    """下载原始文件到本地"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project):
        # 校验参数
        print(project)
        if not all([project]):
            return http.JsonResponse({'code': 1001, 'error': '导出失败'})
        # 查找原始数据
        try:
            a = File_old.objects.get(id=project,user_id=request.user.id)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1004, 'error': "原始表格查询失败"})
        path = a.path
        # 切割路径
        end = '/static/'
        route01 = path[path.rfind(end):]
        context = {
            'path': route01
        }
        return http.JsonResponse({'code': 200, 'error': '成功', 'context': context})


class Article_Download(APIView):
    """文章下载"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project):
        # 校验参数
        print(project)
        if not all([project]):
            return http.JsonResponse({'code': 1001, 'error': '导出失败，请稍后再试'})
        # 查找原始数据
        try:
            a = File_old.objects.get(id=project,user_id=request.user.id)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1004, 'error': "原始表格查询失败"})
        path = a.path
        # 切割路径
        end = '/'
        route01 = path[:path.rfind(end)]
        # 生成word文档的路径
        word_path = route01 + '/' + '文章结果.docx'
        # 切割路径
        end = '/static/'
        route01 = word_path[word_path.rfind(end):]
        context = {
            'path': route01
        }
        return http.JsonResponse({'code': 200, 'error': '成功', 'context': context})


class Survival(APIView):
    """生存分析"""
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        """
        实现生存分析
        :param request:
        :return:
        """
        context = get_s(request, project_id, data_id)
        # 单选加多选
        context["value"] = "dobuleSingle"
        context["begin"] = "end"
        context["name"] = "scfx"
        return render(request, 'index/scfx.html', context=context)

    def post(self, request, project_id, data_id):

        """
        接收传的列名
        :param request:
        :return:
        """

        # df_input,time_column,event_column,group_column=None,model_type=1,sorted=False,intervals=4
        # 获取数据
        json_data = json.loads(request.body.decode())
        # 获取时间列
        column = json_data.get("column")
        # 获取数据类型
        num = json_data.get("number")
        # 获取事件列
        event_column = json_data.get('event_column')
        # 获取分组列
        group_column = json_data.get('group_column')
        # 获取曲线类型
        model_type = json_data.get('model_type')
        # 是否展示置信区间
        conf = json_data.get('sorted')
        # 是否展示删失表
        censor = json_data.get('censor')
        # 是否展示风险表
        risk = json_data.get('risk')
        
        # 获取间隔
        intervals = json_data.get('intervals')
        # print(num, "-=-=num0--==-=-")
        # 获取小数位
        decimal_num = json_data.get("num_select")
        # 获取图片风格
        palette_style = json_data.get('palette_style')
        # 数据筛选
        filter_data = json_data.get("filter")


        try:
            intervals = int(intervals)
        except Exception as e:
            print(e)

        if not group_column:
            group_column = None

        if not intervals:
            intervals = 4

        if not all([str(num), column, event_column, str(model_type)]):
            return HttpResponse(json.dumps({'code': 1001, 'error': '请选择变量'}))
        # 创建一个uuid唯一值
        id = uuid.uuid1()
        id = str(id)
        # 获取传的列名
        if int(num) == 0:
            # 获取数据库的对象
            user = File_old.objects.get(id=project_id,user_id=request.user.id)
            # 获取原文件的路径
            old_file_path = user.path_pa
            print(old_file_path)
        elif int(num) == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 1002, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            old_file_path = re_d.get("df_result")

        try:
            # 打开文件
            # 查询该用户的用户等级
            try:
                grade = Member.objects.get(user_id=request.user.id)
                # 查询当前用户的等级
                member = MemberType.objects.get(id=grade.member_type_id)
            except Exception as e:
                return http.JsonResponse({'code': 1003, 'error': '会员查询失败，请重试'})
            df_r = read(old_file_path)
            df = df_r.iloc[0:int(member.number)]
            
            # 数据筛选
            if filter_data:
                df = filtering_view(df,json_data)
            # 获取路径

            end01 = "/static/"
            end = "/"
            # 本地代码上面
            # end = "\\"
            if int(num) == 0:
                string3 = old_file_path[:old_file_path.rfind(end)]
            else:
                string = old_file_path[:old_file_path.rfind(end)]
                string3 = string[:string.rfind(end)]
            # strint3路径 为用户创的 项目下面 

            # string2 为图片的静态路径
            
            string_tp = string3[string3.rfind(end01):]
            print(string3)
            # file_s 为这方法存图片以及df表的文件夹  绝对路径
            file_s = string3 + "/" + "生存分析" + id
            try:
                # 创建存图片的文件夹
                os.makedirs(file_s)
            except Exception as e:
                print(e)
                return HttpResponse(json.dumps({'code': 1005, 'error': '创建失败'}, cls=NpEncoder))
            
            savepath = file_s + '/'

            
            # 调用生存分析方法
            try:

                # re = R_surv_ana(df_input=df, time='年龄', statu='高血压', groups='LDL分组', methods=0, conf_int =0 , risk_table=0, ncensor_plot=0, x_distance=4, path=savepath,decimal_num=int(decimal_num))
                re = R_surv_ana(df_input=df, time=column, statu=event_column, groups=group_column, methods=int(model_type), conf_int =int(conf) , risk_table=int(risk), ncensor_plot=int(censor), x_distance=int(intervals), path=savepath,decimal_num=int(decimal_num),palette_style=palette_style)
            except Exception as e:
                print(e)
                return http.JsonResponse({'code': 1005, 'error': '分析失败，请选择正确变量'})
            try:
                if re['error']:
                    return http.JsonResponse({'code': 1005, 'error': re['error']})
            except Exception as e:
                print('正常')
            # 提取plt  和str 数据

            str_s = re[1]
            str_s = str_s.replace("\n", "<br/>")

            plts = re[2]

            df_result_dict = re[0]
            

            df_result = df_result_dict

            a = df_result
            # b = df_result_2
            # print(b)
            before = {}
            if list(a):
                # 保存第一列的列名
                before['Listing'] = []
                name = list(a)
                # 循环列名
                for i in range(len(list(a))):
                    before['Listing'].append(name[i])
                    # if i == 8:
                    #     break
                # 添加数据
                before['info'] = []
                for n in range(len(a[name[1]])):
                    dict_x = {}
                    y = 0
                    # 制定第一行 所有列的所有info数据
                    for i in name:
                        i = str(i)
                        # 最多添加八个数据
                        # if y == 8:
                        #     break
                        try:
                            if dict_x[i]:
                                t = str(uuid.uuid1())
                                i = str(i) + t
                                dict_x[i] = str(a.iloc[n][y])
                        except Exception as e:
                            dict_x[i] = str(list(a.iloc[n])[y])
                        y += 1
                    before['info'].append(dict_x)


            # 把得到的图片保存
            # 返回前端 图片的静态路径
            print(plts)
            print(string_tp)
            pltlist = loop_add(plts, "生存分析", string_tp, id)
            # if df_result_dict:
            df_file = file_s + '/' + 'df_result.pkl'

            write(a, df_file)
            browsing_process = {}
            browsing_process['name'] = '生存分析'

            browsing_process['df_result_2'] = df_file
            browsing_process['str_result'] = str_s
            for i in range(len(pltlist)):
                if i == 0:
                    name = 'plt'
                    browsing_process[name] = pltlist[i]
                else:
                    name = 'plt' + str(i + 1)
                    browsing_process[name] = pltlist[i]
                # browsing_process['df_result_2_2'] = df_file2
            try:
                process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                    Max('order'))
                print(process)
                if process:
                    order = int(process['order__max']) + 1
                    print(order)
                    browsing = Browsing_process.objects.create(process_info=browsing_process,
                                                               user_id=request.user.id,
                                                               order=str(order),
                                                               file_old_id=project_id)
                    project = browsing.id

            except Exception as e:
                print(e)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order='1',
                                                           file_old_id=project_id)
                project = browsing.id
            form = [before]
            context = {
                'project': str(project),
                'img': pltlist,
                "form": form,
                "str": str_s,
                "text": "tableDescriptionIImages1"

            }
        except Exception as e:
            print(e)
            return HttpResponse(json.dumps({'code': 2001, 'error': '分析失败,请选择正确的参数'}))
        # return http.JsonResponse({'code':200, 'error':'分析成功','context':context} )
        return HttpResponse(json.dumps({'code': 200, 'error': '分析成功', 'context': context}, cls=NpEncoder))


class REFView(APIView):
    """获取列的数据类型"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project_id):
        """查询相关列的信息"""
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 校验参数
        if not all([features]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if int(number) == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif int(number) == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有保存过最新数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
            # try:
            #     try:
            #         df = pd.read_excel(file_path, nrows=int(member.number))
            #     except Exception as e:
            #         print(e)
            #         df = pd.read_csv(file_path, nrows=int(member.number))
            # except Exception as e:
            #     print(e)
            #     df = pd.read_csv(file_path, nrows=int(member.number), encoding='gbk')
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1005, 'error': '该列数据不正确'})
        # 定义一个
        # 读取列的数据
        df = df.fillna('None')
        name = list(df[str(features)])
        L4 = []
        for x in name:
            if x == 'None':
                continue
            if x not in L4:
                L4.append(x)
        # 返回页面
        return http.JsonResponse({'code': 200, 'error': '分析成功', 'context': L4})


class REFListView(APIView):
    """获取多选列名的下拉框"""
    authentication_classes = [MyBaseAuthentication, ]
    def post(self, request, project_id):
        """查询相关列的信息"""
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 校验参数
        if not all([features]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if int(number) == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif int(number) == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
            # try:
            #     try:
            #         df = pd.read_excel(file_path, nrows=int(member.number))
            #     except Exception as e:
            #         print(e)
            #         df = pd.read_csv(file_path, nrows=int(member.number))
            # except Exception as e:
            #     print(e)
            #     df = pd.read_csv(file_path, nrows=int(member.number),encoding='gbk')
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1005, 'error': '上传文件有误请重新上传'})
        # 定义一个
        # 读取列的数据
        df = df.fillna(' ')
        end_list = {}
        for i in features:
            name = list(df[str(i)])
            L4 = []
            for x in name:
                if x == 'None':
                    continue
                if x not in L4:
                    L4.append(x)
            end_list[i] = L4
        # 返回页面
        return http.JsonResponse({'code': 200, 'error': '分析成功', 'context': end_list})


class ManipulateView(APIView):
    """
    数据操作
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'much'
        return render(request, 'index/data_manipulate.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取数据操作类型
        int_type = json_data.get('method')
        # 获取小数点位数
        decimal_num = json_data.get('decimal_num')
        # 校验参数
        if not all([features, int_type]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
            # try:
            #     try:
            #         df = pd.read_excel(file_path, nrows=int(member.number))
            #     except Exception as e:
            #         print(e)
            #         df = pd.read_csv(file_path, nrows=int(member.number))
            # except Exception as e:
            #     print(e)
            #     df = pd.read_csv(file_path, nrows=int(member.number),encoding='gbk')
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            info = data_manipulate(df_input=df, features=features, type=int(int_type))
        except Exception as e:
            print(e)
            print(1111111111111)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(str(name[i - 1]))
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(c.iloc[num].name): str(c.iloc[num].name)})
                    else:
                        input_o_desc['info'][num][str(list(c)[i - 1])] = str(c.iloc[num][i - 1])
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(str(name[i - 1]))
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(a.iloc[num].name): str(a.iloc[num].name)})
                    else:
                        before['info'][num][str(list(a)[i - 1])] = str(a.iloc[num][i - 1])

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(str(name[i - 1]))
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(d.iloc[num].name): str(d.iloc[num].name)})
                    else:
                        result_o_desc['info'][num][str(list(d)[i - 1])] = str(d.iloc[num][i - 1])

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(str(name[i - 1]))
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(b.iloc[num].name): str(b.iloc[num].name)})
                    else:
                        after['info'][num][str(list(b)[i - 1])] = str(b.iloc[num][i - 1])

        # 将浏览过程添加到数据库
        end = "/"
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/data_manipulate' + str(uuid01))
        except Exception as e:
            print(e)
        print(3333333333333333)
        # 写入数据
        write(info[0], string2 + '/data_manipulate' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/data_manipulate' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/data_manipulate' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/data_manipulate' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/data_manipulate' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '数据操作'
        browsing_process['df_result'] = string2 + '/data_manipulate' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/data_manipulate' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/data_manipulate' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/data_manipulate' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/data_manipulate' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # 如果有图片保存到数据库
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id
        print(4444444444444444444444)
        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")
        data = {
            'form': form,
            'project': project,
            'str_result': str_s,
        }
        return http.JsonResponse({'code': 200, 'error': '分析成功', 'context': data})
        # 返回页面
        # return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))


class Transform(APIView):
    """
    数据转化
    """
    authentication_classes = [MyBaseAuthentication, ]
    def get(self, request, project_id, data_id):
        context = get_s(request, project_id, data_id)
        context['value'] = 'much'
        return render(request, 'index/transform.html', context=context)

    def post(self, request, project_id, data_id):
        # 接收json数据
        json_data = json.loads(request.body.decode())
        # 获取列名列表
        features = json_data.get('features')
        # 获取哪种数据
        number = json_data.get('number')
        # 获取数据操作类型
        int_type = json_data.get('type')
        # 自动转化为定类数据的最大分类数阈值
        decimal_num = json_data.get('num')

        # 校验参数
        if not all([features, int_type]):
            return HttpResponse(json.dumps({'code': 1002, 'error': '变量不能为空'}, cls=NpEncoder))
        # 判断数据是原始数据还是处理后的数据
        file_path = None
        if number == 0:
            # 查询数据库原始文件
            file = File_old.objects.get(id=project_id, user_id=request.user.id)
            file_path = file.path_pa
        elif number == 1:
            try:
                process = Browsing_process.objects.get(user_id=request.user.id, file_old_id=project_id, is_latest=1)
            except Exception as e:
                return HttpResponse(json.dumps({'code': 202, 'error': '你还没有处理过数据，所有没有最新数据噢'}))
            re_d = eval(process.process_info)
            file_path = re_d.get("df_result")
        try:
            # 读取文件
            # 查询该用户的用户等级
            grade = Member.objects.get(user_id=request.user.id)
            # 查询当前用户的等级
            member = MemberType.objects.get(id=grade.member_type_id)
            df_r = read(file_path)
            df = df_r.iloc[0:int(member.number)]
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '上传文件有误请重新上传'})
        # 读取数据
        try:
            if decimal_num:
                decimal_num = int(decimal_num)
            else:
                decimal_num = 5
            info = data_transform(df_input=df, features=features, type=int(int_type), num=decimal_num)
        except Exception as e:
            print(e)
            return http.JsonResponse({'code': 1002, 'error': '分析失败，请选择正确变量'})
        try:
            if info['error']:
                return http.JsonResponse({'code': 1005, 'error': info['error']})
        except Exception as e:
            print('正常')
        # 原始分类数据描述
        c = info[2]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        input_o_desc = {}
        input_o_desc['name'] = '原始分类数据描述'
        if list(c):
            # 保存第一列的列名
            input_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(c)
            # 循环列名
            for i in range(len(list(c)) + 1):
                if i == 0:
                    input_o_desc['Listing'].append('名字')
                else:
                    input_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            input_o_desc['info'] = []
            for num in range(len(c['总数'])):
                for i in range(len(list(c)) + 1):
                    if i == 0:
                        input_o_desc['info'].append({str(c.iloc[num].name): c.iloc[num].name})
                    else:
                        input_o_desc['info'][num][str(list(c)[i - 1])] = c.iloc[num][i - 1]
        # 原始计量数据描述
        a = info[3]
        before = {}
        before['name'] = '原始计量数据描述'
        if list(a):
            # 保存第一列的列名
            before['Listing'] = []
            name = list(a)
            print(name)
            list01 = ['count', 'unique', 'top', 'freq']
            # 循环列名
            for i in range(len(list(a)) + 1):
                if i == 0:
                    before['Listing'].append('名字')
                else:
                    before['Listing'].append(name[i - 1])
            # 添加数据
            before['info'] = []
            for num in range(len(a['总数'])):
                for i in range(len(list(a)) + 1):
                    if i == 0:
                        before['info'].append({str(a.iloc[num].name): a.iloc[num].name})
                    else:
                        before['info'][num][str(list(a)[i - 1])] = a.iloc[num][i - 1]

        # 结果分类数据描述
        d = info[4]
        list02 = ['count', 'categorical', 'highest', 'frequency']
        result_o_desc = {}
        result_o_desc['name'] = '结果分类数据描述'
        if list(d):
            # 保存第一列的列名
            result_o_desc['Listing'] = []
            # 所有列名的列表
            name = list(d)
            # 循环列名
            for i in range(len(list(d)) + 1):
                if i == 0:
                    result_o_desc['Listing'].append('名字')
                else:
                    result_o_desc['Listing'].append(name[i - 1])
            # 添加数据
            result_o_desc['info'] = []
            for num in range(len(d['总数'])):
                for i in range(len(list(d)) + 1):
                    if i == 0:
                        result_o_desc['info'].append({str(d.iloc[num].name): d.iloc[num].name})
                    else:
                        result_o_desc['info'][num][str(list(d)[i - 1])] = d.iloc[num][i - 1]

        # 结果计量数据描述
        b = info[5]
        after = {}
        after['name'] = '结果计量数据描述'
        if list(b):
            # 保存第一列的列名
            after['Listing'] = []
            # 所有列名的列表
            name = list(b)
            # 循环列名
            for i in range(len(list(b)) + 1):
                if i == 0:
                    after['Listing'].append('名字')
                else:
                    after['Listing'].append(name[i - 1])
            # 添加数据
            after['info'] = []
            for num in range(len(b['总数'])):
                for i in range(len(list(b)) + 1):
                    if i == 0:
                        after['info'].append({str(b.iloc[num].name): b.iloc[num].name})
                    else:
                        after['info'][num][str(list(b)[i - 1])] = b.iloc[num][i - 1]

        # 将浏览过程添加到数据库
        end = "/"
        uuid01 = uuid.uuid1()
        if number == 0:
            old_file_path = file_path
            string2 = old_file_path[:old_file_path.rfind(end)]
        else:
            string = file_path[:file_path.rfind(end)]
            string2 = string[:string.rfind(end)]
        try:
            os.mkdir(string2 + '/data_transform' + str(uuid01))
        except Exception as e:
            print(e)
        
        # 写入数据
        write(info[0], string2 + '/data_transform' + str(uuid01) + '/' + 'df_result.pkl')
        write(info[2], string2 + '/data_transform' + str(uuid01) + '/' + 'input_o_desc.pkl')
        write(info[3], string2 + '/data_transform' + str(uuid01) + '/' + 'input_n_desc.pkl')
        write(info[4], string2 + '/data_transform' + str(uuid01) + '/' + 'result_o_desc.pkl')
        write(info[5], string2 + '/data_transform' + str(uuid01) + '/' + 'result_n_desc.pkl')

        # 写进数据库
        browsing_process = {}
        browsing_process['name'] = '数据转换'
        browsing_process['df_result'] = string2 + '/data_transform' + str(uuid01) + '/' + 'df_result.pkl'
        browsing_process['str_result'] = info[1]
        browsing_process['input_n_desc'] = string2 + '/data_transform' + str(uuid01) + '/' + 'input_n_desc.pkl'
        browsing_process['input_o_desc'] = string2 + '/data_transform' + str(uuid01) + '/' + 'input_o_desc.pkl'
        browsing_process['result_o_desc'] = string2 + '/data_transform' + str(uuid01) + '/' + 'result_o_desc.pkl'
        browsing_process['result_n_desc'] = string2 + '/data_transform' + str(uuid01) + '/' + 'result_n_desc.pkl'
        # 如果有图片保存到数据库
        # # 查询所有之前的顺序
        try:
            process = Browsing_process.objects.filter(user_id=request.user.id, file_old_id=project_id).aggregate(
                Max('order'))
            print(process)
            if process:
                order = int(process['order__max']) + 1
                print(order)
                browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id,
                                                           order=order,
                                                           file_old_id=project_id)
                project = browsing.id

        except Exception as e:
            print(e)
            browsing = Browsing_process.objects.create(process_info=browsing_process, user_id=request.user.id, order=1,
                                                       file_old_id=project_id)
            project = browsing.id

        form = [before, after, input_o_desc, result_o_desc]
        str_s = info[1]
        str_s = str_s.replace("\n", "<br/>")

        data = {
            'form': form,
            'project': project,
            'str_result': str_s
        }

        # 返回页面
        return HttpResponse(json.dumps({'code': '200', 'error': '分析成功', 'context': data}, cls=NpEncoder))































































